"""Document processing and extraction."""
import io
from pathlib import Path
from typing import Optional
import pypdfium2 as pdfium
import pandas as pd
from docx import Document as DocxDocument
from pydantic import BaseModel
from loguru import logger


class ProcessedDocument(BaseModel):
    """Processed document with extracted content."""
    filename: str
    content: str
    doc_type: str
    num_pages: Optional[int] = None
    metadata: dict = {}


class DocumentProcessor:
    """Process various document types."""
    
    @staticmethod
    async def process_file(file_path: Path, file_content: Optional[bytes] = None) -> ProcessedDocument:
        """Process a file and extract content."""
        suffix = file_path.suffix.lower()
        filename = file_path.name
        
        try:
            if suffix == ".pdf":
                return await DocumentProcessor._process_pdf(filename, file_path, file_content)
            elif suffix in [".txt", ".md"]:
                return await DocumentProcessor._process_text(filename, file_path, file_content)
            elif suffix in [".csv"]:
                return await DocumentProcessor._process_csv(filename, file_path, file_content)
            elif suffix in [".xlsx", ".xls"]:
                return await DocumentProcessor._process_excel(filename, file_path, file_content)
            elif suffix in [".docx"]:
                return await DocumentProcessor._process_docx(filename, file_path, file_content)
            else:
                raise ValueError(f"Unsupported file type: {suffix}")
        except Exception as e:
            logger.error(f"Error processing {filename}: {e}")
            raise
    
    @staticmethod
    async def _process_pdf(filename: str, file_path: Path, content: Optional[bytes]) -> ProcessedDocument:
        """Extract text from PDF."""
        if content:
            pdf = pdfium.PdfDocument(io.BytesIO(content))
        else:
            pdf = pdfium.PdfDocument(file_path)
        
        text_parts = []
        for i, page in enumerate(pdf):
            textpage = page.get_textpage()
            text = textpage.get_text_range()
            text_parts.append(f"[Page {i+1}]\n{text}\n")
        
        full_text = "\n".join(text_parts)
        
        return ProcessedDocument(
            filename=filename,
            content=full_text,
            doc_type="pdf",
            num_pages=len(pdf),
            metadata={"page_count": len(pdf)}
        )
    
    @staticmethod
    async def _process_text(filename: str, file_path: Path, content: Optional[bytes]) -> ProcessedDocument:
        """Process plain text file."""
        if content:
            text = content.decode("utf-8")
        else:
            text = file_path.read_text(encoding="utf-8")
        
        return ProcessedDocument(
            filename=filename,
            content=text,
            doc_type="text"
        )
    
    @staticmethod
    async def _process_csv(filename: str, file_path: Path, content: Optional[bytes]) -> ProcessedDocument:
        """Process CSV file."""
        if content:
            df = pd.read_csv(io.BytesIO(content))
        else:
            df = pd.read_csv(file_path)
        
        # Create text representation
        text_parts = [
            f"CSV File: {filename}",
            f"Rows: {len(df)}, Columns: {len(df.columns)}",
            f"\nColumn Names:\n{', '.join(df.columns.tolist())}",
            f"\nFirst 5 rows:\n{df.head().to_string()}",
            f"\nData Summary:\n{df.describe().to_string()}"
        ]
        
        content_text = "\n\n".join(text_parts)
        
        return ProcessedDocument(
            filename=filename,
            content=content_text,
            doc_type="csv",
            metadata={
                "rows": len(df),
                "columns": len(df.columns),
                "column_names": df.columns.tolist()
            }
        )
    
    @staticmethod
    async def _process_excel(filename: str, file_path: Path, content: Optional[bytes]) -> ProcessedDocument:
        """Process Excel file."""
        if content:
            df = pd.read_excel(io.BytesIO(content))
        else:
            df = pd.read_excel(file_path)
        
        text_parts = [
            f"Excel File: {filename}",
            f"Rows: {len(df)}, Columns: {len(df.columns)}",
            f"\nColumn Names:\n{', '.join(df.columns.tolist())}",
            f"\nFirst 5 rows:\n{df.head().to_string()}",
            f"\nData Summary:\n{df.describe().to_string()}"
        ]
        
        content_text = "\n\n".join(text_parts)
        
        return ProcessedDocument(
            filename=filename,
            content=content_text,
            doc_type="excel",
            metadata={
                "rows": len(df),
                "columns": len(df.columns),
                "column_names": df.columns.tolist()
            }
        )
    
    @staticmethod
    async def _process_docx(filename: str, file_path: Path, content: Optional[bytes]) -> ProcessedDocument:
        """Process Word document."""
        if content:
            doc = DocxDocument(io.BytesIO(content))
        else:
            doc = DocxDocument(file_path)
        
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        full_text = "\n\n".join(text_parts)
        
        return ProcessedDocument(
            filename=filename,
            content=full_text,
            doc_type="docx",
            metadata={"paragraph_count": len(doc.paragraphs)}
        )